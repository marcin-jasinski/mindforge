"""
DOCX document parser.

Extracts text, document properties (title, author, subject), and embedded
images from DOCX files using ``python-docx``.
"""

from __future__ import annotations

import io
from typing import Any

try:
    import docx
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
except ImportError as exc:
    raise ImportError(
        "python-docx is required for DocxParser. "
        "Install it via: pip install python-docx"
    ) from exc

from mindforge.domain.models import BlockType, ContentBlock, ParsedDocument
from mindforge.infrastructure.parsing.registry import ParseError


DEFAULT_MAX_IMAGE_BYTES = 10 * 1024 * 1024   # 10 MB per image
DEFAULT_MAX_TOTAL_IMAGE_BYTES = 50 * 1024 * 1024  # 50 MB across all images


class DocxParser:
    """
    Parse DOCX documents using python-docx.

    Extracts:
    - Full paragraph text (in document order)
    - Core document properties: ``title``, ``author``, ``subject``
    - Embedded images as raw bytes (for vision-model analysis)
    """

    def __init__(
        self,
        max_image_bytes: int = DEFAULT_MAX_IMAGE_BYTES,
        max_total_image_bytes: int = DEFAULT_MAX_TOTAL_IMAGE_BYTES,
    ) -> None:
        self._max_image_bytes = max_image_bytes
        self._max_total_image_bytes = max_total_image_bytes

    def parse(self, raw_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
        except Exception as exc:
            raise ParseError(f"Failed to open DOCX: {exc}") from exc

        return self._extract(doc, filename)

    def _extract(self, doc: DocxDocument, filename: str) -> ParsedDocument:
        # -- Metadata -------------------------------------------------------
        metadata: dict[str, Any] = {}
        try:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
        except Exception:
            pass  # Core properties are optional

        # -- Paragraphs -----------------------------------------------------
        content_blocks: list[ContentBlock] = []
        paragraphs_text: list[str] = []
        position = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            paragraphs_text.append(text)

            style_name = para.style.name if para.style else ""
            is_heading = style_name.lower().startswith("heading")

            content_blocks.append(
                ContentBlock(
                    block_type=BlockType.TEXT,
                    content=text,
                    position=position,
                    metadata={"style": style_name, "is_heading": is_heading},
                )
            )
            position += 1

        # -- Embedded images ------------------------------------------------
        embedded_images: list[bytes] = []
        total_image_bytes = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_bytes: bytes = image_part.blob
                    content_type: str = image_part.content_type or "image/png"
                    if not image_bytes:
                        continue
                    if len(image_bytes) > self._max_image_bytes:
                        # Skip individual oversized images (non-fatal)
                        continue
                    if total_image_bytes + len(image_bytes) > self._max_total_image_bytes:
                        # Cumulative cap reached; stop collecting images
                        break
                    total_image_bytes += len(image_bytes)
                    embedded_images.append(image_bytes)
                    content_blocks.append(
                        ContentBlock(
                            block_type=BlockType.IMAGE,
                            content="",
                            media_ref=rel.target_ref,
                            media_type=content_type,
                            position=position,
                        )
                    )
                    position += 1
                except Exception:
                    pass  # Non-fatal

        full_text = "\n".join(paragraphs_text)

        # Extract first heading for title resolution if not in core props
        if "title" not in metadata:
            for block in content_blocks:
                if block.metadata.get("is_heading"):
                    metadata.setdefault("first_heading", block.content)
                    break

        return ParsedDocument(
            text_content=full_text,
            metadata=metadata,
            content_blocks=content_blocks,
            embedded_images=embedded_images,
        )
